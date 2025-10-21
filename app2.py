import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import io
import tempfile
from fpdf import FPDF
from docx import Document
from tabulate import tabulate
import os

st.set_page_config(page_title="Dashboard Marketplace & Incubateur", layout="wide")
st.title("Dashboard Marketplace & Incubateur")

# --- Fonction utilitaire pour nettoyer le texte pour PDF ---
def clean_text(text):
    if pd.isna(text):
        return ""
    return str(text).replace("\n", " ").replace("\r", " ").replace("\t", " ").strip()

# --- Fonction pour lire CSV/XLSX ---
def read_file_safe(uploaded_file, expected_columns=None):
    if uploaded_file is None or uploaded_file.size == 0:
        st.error(f"Le fichier {uploaded_file.name if uploaded_file else 'inconnu'} est vide !")
        return pd.DataFrame()

    content = uploaded_file.read()
    buffer = io.BytesIO(content)
    df = None

    if uploaded_file.name.endswith(".csv"):
        encodings = ["utf-8", "utf-8-sig", "ISO-8859-1"]
        for enc in encodings:
            try:
                buffer.seek(0)
                df = pd.read_csv(buffer, encoding=enc, engine="python", on_bad_lines='skip')
                break
            except Exception:
                df = None
        if df is None:
            st.error(f"Impossible de lire le fichier {uploaded_file.name} avec tous les encodages")
            return pd.DataFrame()
    elif uploaded_file.name.endswith(".xlsx"):
        buffer.seek(0)
        try:
            df = pd.read_excel(buffer, engine="openpyxl")
        except Exception as e:
            st.error(f"Impossible de lire le fichier {uploaded_file.name} : {e}")
            return pd.DataFrame()
    else:
        st.error(f"Format de fichier non supporté : {uploaded_file.name}")
        return pd.DataFrame()

    df.columns = df.columns.str.strip()
    if expected_columns:
        missing_cols = [c for c in expected_columns if c not in df.columns]
        if missing_cols:
            st.error(f"Colonnes manquantes dans {uploaded_file.name} : {missing_cols}")
    return df

# --- Upload fichiers ---
st.sidebar.header("Uploader les fichiers")
file_users = st.sidebar.file_uploader("Noms persos", type=["csv","xlsx"])
file_entreprises = st.sidebar.file_uploader("Entreprises données", type=["csv","xlsx"])
file_mises_relation = st.sidebar.file_uploader("Historique des mises en relation", type=["csv","xlsx"])
file_base_globale = st.sidebar.file_uploader("Base globale projet", type=["csv","xlsx"])

# --- Colonnes attendues ---
cols_users = ["#Id", "Prénom", "Nom", "Inscrit depuis le", "Statut", "ID Unique", "Date de dernière connexion"]
cols_entreprises = ["Id", "Nom", "Date de création", "Date d'ouverture", "Incubateurs", "À propos", "Missions",
                    "Adresse", "Ville", "Code postal", "Téléphone", "Email", "Effectifs", "Linkedin", "Site web",
                    "Équipe", "Statut"]
cols_mises = ["Utilisateur","goBetween","Statut des mises en relation à date","Dates simples",
              "Demande de mise en relation","RDV réalisés","Taux de conversion goBetween",
              "Taux de conversion RDV réalisé","Go between validé","Go between refusé","Rdv non réalisé"]
cols_globale = ["Name","Nom","Projet","CAR/SUM (territorial)","Incubateur territorial","Statut d'incubation",
                "Poste et/ou fonction","Profil personnel Le Club","Profil sociétés Le Club",
                "Partenaires Marketplace","Date dernière connexion Le Club"]

# --- Lecture sécurisée ---
df_users = read_file_safe(file_users, expected_columns=cols_users)
df_entreprises = read_file_safe(file_entreprises, expected_columns=cols_entreprises)
df_mises = read_file_safe(file_mises_relation, expected_columns=cols_mises)
df_globale = read_file_safe(file_base_globale, expected_columns=cols_globale)

# --- Vérification fichiers valides ---
if not df_users.empty and not df_entreprises.empty and not df_mises.empty and not df_globale.empty:

    # --- Datas globales ---
    st.header("Datas globales")
    demandes_total = len(df_mises)
    profils_total = len(df_users)
    
    today = datetime.today()
    month_ago = today - timedelta(days=30)
    df_users["Date de dernière connexion"] = pd.to_datetime(df_users["Date de dernière connexion"], dayfirst=True, errors="coerce")
    profils_connectes = df_users[df_users["Date de dernière connexion"] >= month_ago].shape[0]
    
    st.metric("Demandes de mise en relation", demandes_total)
    st.metric("Profils créés", profils_total)
    st.metric("Profils connectés sur le mois", profils_connectes)

    # --- Marketplace ---
    st.header("Marketplace")
    df_mises["Go between validé"] = df_mises["Go between validé"].astype(str).str.strip()
    df_mises["Rdv non réalisé"] = pd.to_numeric(df_mises["Rdv non réalisé"], errors="coerce").fillna(0)
    df_mises["RDV réalisés"] = pd.to_numeric(df_mises["RDV réalisés"], errors="coerce").fillna(0)

    go_between_valides = (df_mises["Go between validé"].str.lower() == "oui").sum()
    rdv_realises = df_mises["RDV réalisés"].sum()
    rdv_non_realises = df_mises["Rdv non réalisé"].sum()
    
    st.subheader("Totaux")
    st.metric("Go Between validés", go_between_valides)
    st.metric("RDV réalisés", rdv_realises)
    st.metric("RDV non réalisés", rdv_non_realises)

    st.subheader("Statut des mises en relation")
    statut_mises_table = df_mises["Statut des mises en relation à date"].astype(str).str.strip().value_counts()
    st.table(statut_mises_table)

    taux_go_between = pd.to_numeric(df_mises["Taux de conversion goBetween"], errors="coerce").mean()
    taux_rdv = pd.to_numeric(df_mises["Taux de conversion RDV réalisé"], errors="coerce").mean()
    st.metric("Taux de conversion Go Between (%)", round(taux_go_between, 2))
    st.metric("Taux de conversion RDV réalisés (%)", round(taux_rdv, 2))

    df_mises["Dates simples"] = pd.to_datetime(df_mises["Dates simples"], format="%Y-%m", errors="coerce")
    df_mises["Trimestre"] = df_mises["Dates simples"].dt.to_period("Q")
    trimestriel = df_mises.groupby("Trimestre").size()
    st.subheader("Répartition trimestrielle des demandes")
    st.table(trimestriel)

    # --- Profils persos & Sociétés ---
    st.header("Profils persos & Sociétés")
    st.metric("Nombre total d'entrepreneurs", len(df_globale))
    st.metric("Total profils persos", len(df_users))

    st.subheader("Statut profils persos")
    statut_profil_persos = df_users["Statut"].astype(str).str.strip().value_counts()
    st.table(statut_profil_persos)

    st.subheader("Statut profils sociétés")
    statut_profil_societes = df_entreprises["Statut"].astype(str).str.strip().value_counts()
    st.table(statut_profil_societes)

    # --- Complétion des profils ---
    st.header("Complétion des profils")
    st.subheader("Vue globale")
    profil_persos_counts = df_globale["Profil personnel Le Club"].astype(str).str.strip().value_counts()
    profil_societes_counts = df_globale["Profil sociétés Le Club"].astype(str).str.strip().value_counts()
    st.table(profil_persos_counts)
    st.table(profil_societes_counts)

    st.subheader("Par CAR/SUM")
    car_sum_counts = df_globale.groupby("CAR/SUM (territorial)")["Profil sociétés Le Club"].value_counts()
    st.table(car_sum_counts)

    st.subheader("Par Incubateur territorial")
    incubateur_counts = df_globale.groupby("Incubateur territorial")["Profil sociétés Le Club"].value_counts()
    st.table(incubateur_counts)

    st.subheader("% de complétion sur les profils incubation individuelle")
    incubation_indiv = df_globale[df_globale["Statut d'incubation"].astype(str).str.strip() == "Incubation individuelle"]
    incubation_counts = (incubation_indiv["Profil sociétés Le Club"]
                         .astype(str)
                         .str.strip()
                         .value_counts(normalize=True)
                         .mul(100)
                         .round(2))
    st.table(incubation_counts)

    # --- Génération PDF ---
    def generate_pdf():
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        # Ajouter une police Unicode si nécessaire
        font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        if os.path.exists(font_path):
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=10)
        else:
            pdf.set_font("Arial", size=10)

        def write_table(series, title):
            pdf.multi_cell(0, 5, clean_text(title))
            table_text = tabulate(series.reset_index().values, headers=[series.index.name or "Catégorie","Nombre"], tablefmt="grid")
            for line in table_text.split("\n"):
                pdf.multi_cell(0, 5, clean_text(line))
            pdf.ln(5)

        pdf.multi_cell(0, 5, f"Demandes de mise en relation: {demandes_total}")
        pdf.multi_cell(0, 5, f"Profils créés: {profils_total}")
        pdf.multi_cell(0, 5, f"Profils connectés sur le mois: {profils_connectes}")
        pdf.ln(5)

        write_table(statut_mises_table, "Statut des mises en relation")
        write_table(statut_profil_persos, "Statut profils persos")
        write_table(statut_profil_societes, "Statut profils sociétés")
        write_table(profil_persos_counts, "Complétion profils persos")
        write_table(profil_societes_counts, "Complétion profils sociétés")
        write_table(car_sum_counts, "Complétion par CAR/SUM")
        write_table(incubateur_counts, "Complétion par Incubateur territorial")
        write_table(incubation_counts, "Complétion sur Incubation individuelle (%)")

        return pdf

    if st.button("Télécharger PDF"):
        pdf = generate_pdf()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf.output(tmp.name)
            tmp.seek(0)
            st.download_button("Télécharger PDF", tmp.name, file_name="KPIs.pdf", mime="application/pdf")

    # --- Génération DOCX ---
    def generate_docx():
        doc = Document()
        doc.add_heading("Dashboard Marketplace & Incubateur", level=1)

        def write_table_doc(series, title):
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = series.index.name or "Catégorie"
            hdr_cells[1].text = "Nombre"
            for idx, val in series.items():
                row_cells = table.add_row().cells
                row_cells[0].text = clean_text(idx)
                row_cells[1].text = str(val)
            doc.add_paragraph("\n")

        doc.add_paragraph(f"Demandes de mise en relation: {demandes_total}")
        doc.add_paragraph(f"Profils créés: {profils_total}")
        doc.add_paragraph(f"Profils connectés sur le mois: {profils_connectes}")
        write_table_doc(statut_mises_table, "Statut des mises en relation")
        write_table_doc(statut_profil_persos, "Statut profils persos")
        write_table_doc(statut_profil_societes, "Statut profils sociétés")
        write_table_doc(profil_persos_counts, "Complétion profils persos")
        write_table_doc(profil_societes_counts, "Complétion profils sociétés")
        write_table_doc(car_sum_counts, "Complétion par CAR/SUM")
        write_table_doc(incubateur_counts, "Complétion par Incubateur territorial")
        write_table_doc(incubation_counts, "Complétion sur Incubation individuelle (%)")

        return doc

    if st.button("Télécharger DOCX"):
        doc = generate_docx()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button("Télécharger DOCX", tmp.name, file_name="KPIs.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.info("Veuillez uploader tous les fichiers correctement pour générer les KPIs.")
