import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import io
from docx import Document

# -----------------------------
# Thème global
# -----------------------------
st.markdown("""
<style>
.stApp {
    background-color: #1d2732;
    color: white;
    font-family: 'Inter', sans-serif;
}

h1, h2, h3 {
    color: #3ecdd1;
    font-weight: 600;
}

section[data-testid="stSidebar"] {
    background-color: #232f3c;
    color: white;
}

.stButton > button {
    background-color: #3ecdd1;
    color: #1d2732;
    border: none;
    border-radius: 6px;
    padding: 0.4em 1.2em;
    font-weight: 600;
    transition: all 0.2s ease;
}
.stButton > button:hover {
    background-color: #35b9c1;
    color: white;
}

.stMetric {
    background-color: #232f3c;
    border-radius: 8px;
    padding: 0.8em;
}
.stMetric label {
    color: #3ecdd1 !important;
}

.stTable {
    background-color: #232f3c;
    border-radius: 8px;
    padding: 0.5em;
}

hr {
    border: none;
    border-top: 1px solid #3ecdd1;
    margin: 1.5em 0;
}
</style>
""", unsafe_allow_html=True)

# -----------------------------
# Sidebar logos et uploads
# -----------------------------
st.sidebar.image("logo1.png", use_container_width=True)
st.sidebar.markdown("<br>", unsafe_allow_html=True)
st.sidebar.image("logo2.png", use_container_width=True)

st.sidebar.header("Uploader les fichiers")
file_users = st.sidebar.file_uploader("Noms persos", type=["csv","xlsx"])
file_entreprises = st.sidebar.file_uploader("Entreprises données", type=["csv","xlsx"])
file_mises_relation = st.sidebar.file_uploader("Historique des mises en relation", type=["csv","xlsx"])
file_base_globale = st.sidebar.file_uploader("Base globale projet", type=["csv","xlsx"])

st.title("Dashboard Marketplace & Incubateur")

# -----------------------------
# Lecture sécurisée
# -----------------------------
def read_file_safe(uploaded_file, expected_columns=None):
    if uploaded_file is None or uploaded_file.size == 0:
        st.error(f"Le fichier {uploaded_file.name if uploaded_file else 'inconnu'} est vide !")
        return pd.DataFrame()

    content = uploaded_file.read()
    buffer = io.BytesIO(content)

    df = None
    if uploaded_file.name.endswith(".csv"):
        encodings = ["utf-8", "utf-8-sig", "ISO-8859-1"]
        for enc in encodings:
            try:
                buffer.seek(0)
                df = pd.read_csv(buffer, encoding=enc, engine="python")
                break
            except Exception:
                df = None
        if df is None:
            st.error(f"Impossible de lire le fichier {uploaded_file.name} avec tous les encodages")
            return pd.DataFrame()
    elif uploaded_file.name.endswith(".xlsx"):
        buffer.seek(0)
        try:
            df = pd.read_excel(buffer)
        except Exception as e:
            st.error(f"Impossible de lire le fichier {uploaded_file.name} : {e}")
            return pd.DataFrame()
    else:
        st.error(f"Format de fichier non supporté : {uploaded_file.name}")
        return pd.DataFrame()

    df.columns = df.columns.str.strip()
    if expected_columns:
        missing_cols = [c for c in expected_columns if c not in df.columns]
        if missing_cols:
            st.error(f"Colonnes manquantes dans {uploaded_file.name} : {missing_cols}")

    return df

# -----------------------------
# Colonnes attendues
# -----------------------------
cols_users = ["#Id", "Prénom", "Nom", "Inscrit depuis le", "Statut", "ID Unique", "Date de dernière connexion"]
cols_entreprises = ["Id", "Nom", "Date de création", "Date d'ouverture", "Incubateurs", "À propos", "Missions",
                    "Adresse", "Ville", "Code postal", "Téléphone", "Email", "Effectifs", "Linkedin", "Site web",
                    "Équipe", "Statut"]
cols_mises = ["Utilisateur","goBetween","Statut des mises en relation à date","Dates simples",
              "Demande de mise en relation","RDV réalisés","Taux de conversion goBetween",
              "Taux de conversion RDV réalisé","Go between validé","Go between refusé","Rdv non réalisé"]
cols_globale = ["Name","Nom","Projet","CAR/SUM (territorial)","Incubateur territorial","Statut d'incubation",
                "Poste et/ou fonction","Profil personnel Le Club","Profil sociétés Le Club",
                "Partenaires Marketplace","Date dernière connexion Le Club"]

# -----------------------------
# Dashboard principal avec spinner
# -----------------------------
with st.spinner("Chargement du dashboard..."):
    df_users = read_file_safe(file_users, expected_columns=cols_users)
    df_entreprises = read_file_safe(file_entreprises, expected_columns=cols_entreprises)
    df_mises = read_file_safe(file_mises_relation, expected_columns=cols_mises)
    df_globale = read_file_safe(file_base_globale, expected_columns=cols_globale)

    if not df_users.empty and not df_entreprises.empty and not df_mises.empty and not df_globale.empty:

        today = datetime.today()
        month_ago = today - timedelta(days=30)
        df_users["Date de dernière connexion"] = pd.to_datetime(df_users["Date de dernière connexion"], dayfirst=True, errors="coerce")

        # --- Datas globales ---
        st.header("Datas globales")
        demandes_total = len(df_mises)
        profils_total = len(df_users)
        profils_connectes = df_users[df_users["Date de dernière connexion"] >= month_ago].shape[0]

        st.metric("Demandes de mise en relation", demandes_total)
        st.metric("Profils créés", profils_total)
        st.metric("Profils connectés sur le mois", profils_connectes)

        # --- Marketplace ---
        st.header("Marketplace")
        go_between_valides = (df_mises["Go between validé"] == "Oui").sum()
        rdv_realises = df_mises["RDV réalisés"].sum()
        rdv_non_realises = df_mises["Rdv non réalisé"].sum()
        taux_go_between = round(pd.to_numeric(df_mises["Taux de conversion goBetween"], errors="coerce").mean(), 2)
        taux_rdv = round(pd.to_numeric(df_mises["Taux de conversion RDV réalisé"], errors="coerce").mean(), 2)

        st.metric("Go Between validés", go_between_valides)
        st.metric("RDV réalisés", rdv_realises)
        st.metric("RDV non réalisés", rdv_non_realises)
        st.metric("Taux de conversion Go Between (%)", taux_go_between)
        st.metric("Taux de conversion RDV réalisés (%)", taux_rdv)

        # Totaux trimestriels
        df_mises["Dates simples"] = pd.to_datetime(df_mises["Dates simples"], format="%Y-%m", errors="coerce")
        df_mises["Trimestre"] = df_mises["Dates simples"].dt.to_period("Q")
        trimestriel = df_mises.groupby("Trimestre").size()
        st.header("Totaux trimestriels")
        st.table(trimestriel)

        # --- Profils persos & Sociétés ---
        st.header("Profils persos & Sociétés")
        st.metric("Nombre total d'entrepreneurs", len(df_globale))
        st.metric("Total profils persos", len(df_users))
        st.table(df_users["Statut"].value_counts())
        st.table(df_entreprises["Statut"].value_counts())

        # --- Complétion des profils ---
        st.header("Complétion des profils")
        st.table(df_globale["Profil personnel Le Club"].value_counts())
        st.table(df_globale["Profil sociétés Le Club"].value_counts())

        # --- Génération DOCX simplifiée (datas finales uniquement) ---
        def generate_docx_metrics(df_users, df_entreprises, df_mises, df_globale):
            doc = Document()
            doc.add_heading("Dashboard Marketplace & Incubateur - Extract", 0)

            # Logo
            try:
                doc.add_picture("logo1.png")
            except:
                pass

            today = datetime.today()
            month_ago = today - timedelta(days=30)
            doc.add_paragraph(f"Généré le {today.strftime('%d/%m/%Y')}")

            # Datas globales
            doc.add_paragraph(f"Demandes de mise en relation: {len(df_mises)}")
            doc.add_paragraph(f"Profils créés: {len(df_users)}")
            doc.add_paragraph(f"Profils connectés sur le mois: {df_users[df_users['Date de dernière connexion'] >= month_ago].shape[0]}")

            # Marketplace
            go_between_valides = (df_mises["Go between validé"] == "Oui").sum()
            rdv_realises = df_mises["RDV réalisés"].sum()
            rdv_non_realises = df_mises["Rdv non réalisé"].sum()
            taux_go_between = round(pd.to_numeric(df_mises["Taux de conversion goBetween"], errors="coerce").mean(), 2)
            taux_rdv = round(pd.to_numeric(df_mises["Taux de conversion RDV réalisé"], errors="coerce").mean(), 2)

            doc.add_paragraph(f"Go Between validés: {go_between_valides}")
            doc.add_paragraph(f"RDV réalisés: {rdv_realises}")
            doc.add_paragraph(f"RDV non réalisés: {rdv_non_realises}")
            doc.add_paragraph(f"Taux de conversion Go Between (%): {taux_go_between}")
            doc.add_paragraph(f"Taux de conversion RDV réalisés (%): {taux_rdv}")

            # Totaux trimestriels
            df_mises["Dates simples"] = pd.to_datetime(df_mises["Dates simples"], format="%Y-%m", errors="coerce")
            df_mises["Trimestre"] = df_mises["Dates simples"].dt.to_period("Q")
            trimestriel = df_mises.groupby("Trimestre").size()
            doc.add_heading("Totaux trimestriels", level=1)
            for trimestre, total in trimestriel.items():
                doc.add_paragraph(f"{trimestre}: {total}")

            # Sauvegarde DOCX
            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return stream

        docx_data = generate_docx_metrics(df_users, df_entreprises, df_mises, df_globale)
        st.download_button(
            label="Télécharger l'extract final en DOCX",
            data=docx_data,
            file_name=f"dashboard_extract_{today.strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.info("Veuillez uploader tous les fichiers pour générer le dashboard.")
